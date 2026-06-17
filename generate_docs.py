#!/usr/bin/env python3
"""
Generate DOCX, PDF, and Excel Gantt chart from MVP roadmap
"""

import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Output directory
OUTPUT_DIR = "/Users/dhiazfathra/Documents/GitHub/dhiazfathra/figma-mcp/docs/superpowers/specs"

# Sprint data
SPRINTS = [
    # Backend sprints
    {"sprint": "B-1", "week": 1, "start": "2026-06-22", "duration": 7, "focus": "Foundation", "deliverables": "Project setup, Auth APIs, User model, DB schema", "team": "Backend", "color": "#4472C4"},
    {"sprint": "B-2", "week": 2, "start": "2026-06-29", "duration": 7, "focus": "Core APIs", "deliverables": "Brand, Booking, Promotion APIs", "team": "Backend", "color": "#4472C4"},
    {"sprint": "B-3", "week": 3, "start": "2026-07-06", "duration": 7, "focus": "Extended APIs", "deliverables": "Notification, Profile, Search APIs", "team": "Backend", "color": "#4472C4"},
    {"sprint": "B-4", "week": 4, "start": "2026-07-13", "duration": 7, "focus": "Domain APIs", "deliverables": "Treatment, Clinic, Skin Analysis APIs", "team": "Backend", "color": "#4472C4"},
    {"sprint": "B-5", "week": 5, "start": "2026-07-20", "duration": 7, "focus": "Phase 2 APIs", "deliverables": "Healthy Food, Shop APIs (Cart, Order)", "team": "Backend", "color": "#4472C4"},
    {"sprint": "B-6", "week": 6, "start": "2026-07-27", "duration": 7, "focus": "Phase 2 APIs", "deliverables": "Treatment Transaction, Skin Journey APIs", "team": "Backend", "color": "#4472C4"},
    {"sprint": "B-7", "week": 7, "start": "2026-08-03", "duration": 7, "focus": "CMS", "deliverables": "Dashboard, Content management", "team": "Backend", "color": "#4472C4"},
    {"sprint": "B-8", "week": 8, "start": "2026-08-10", "duration": 7, "focus": "Polish", "deliverables": "Documentation, Optimization, Security", "team": "Backend", "color": "#4472C4"},
    {"sprint": "B-BUF-1", "week": 9, "start": "2026-08-17", "duration": 7, "focus": "Buffer", "deliverables": "Integration testing, Bug fixes, API documentation", "team": "Backend", "color": "#70AD47"},
    {"sprint": "B-BUF-2", "week": 10, "start": "2026-08-24", "duration": 7, "focus": "Buffer", "deliverables": "Performance optimization, Security audit", "team": "Backend", "color": "#70AD47"},
    
    # Mobile sprints
    {"sprint": "M-1", "week": 3, "start": "2026-07-06", "duration": 7, "focus": "Setup & Auth", "deliverables": "Flutter project, Login screen, OTP verification", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-2", "week": 4, "start": "2026-07-13", "duration": 7, "focus": "Auth Complete", "deliverables": "Register, Forgot Password, Onboarding tutorial", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "BUF-1", "week": 5, "start": "2026-07-20", "duration": 7, "focus": "Buffer", "deliverables": "Auth integration testing, Bug fixes, Documentation", "team": "Mobile", "color": "#70AD47"},
    {"sprint": "M-3", "week": 6, "start": "2026-07-27", "duration": 7, "focus": "Homepage", "deliverables": "Non-login homepage, Logged-in homepage", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-4", "week": 7, "start": "2026-08-03", "duration": 7, "focus": "Homepage Complete", "deliverables": "Search functionality, Brand carousel", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-5", "week": 8, "start": "2026-08-10", "duration": 7, "focus": "Brands", "deliverables": "Brand listing, Brand detail", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-6", "week": 9, "start": "2026-08-17", "duration": 7, "focus": "Brands Complete", "deliverables": "Clinic page, Treatment page", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-7", "week": 10, "start": "2026-08-24", "duration": 7, "focus": "Booking (Part 1)", "deliverables": "Clinic selection, Treatment selection, Date picker", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-8", "week": 11, "start": "2026-08-31", "duration": 7, "focus": "Booking (Part 2)", "deliverables": "Time picker, Booking confirmation, Payment", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-9", "week": 12, "start": "2026-09-07", "duration": 7, "focus": "Booking Complete", "deliverables": "Booking status, Push notifications", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "BUF-2", "week": 13, "start": "2026-09-14", "duration": 7, "focus": "Buffer", "deliverables": "Booking integration testing, Payment flow validation, Bug fixes", "team": "Mobile", "color": "#70AD47"},
    {"sprint": "M-10", "week": 14, "start": "2026-09-21", "duration": 7, "focus": "Profile", "deliverables": "Profile page, Edit profile, Address management", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-11", "week": 15, "start": "2026-09-28", "duration": 7, "focus": "Profile Complete", "deliverables": "Change phone, Change password, Face ID, Privacy", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-12", "week": 16, "start": "2026-10-05", "duration": 7, "focus": "Promotions", "deliverables": "Promotion listing, Promotion detail", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-13", "week": 17, "start": "2026-10-12", "duration": 7, "focus": "History & Skin", "deliverables": "Booking history, Skin analysis questionnaire", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-14", "week": 18, "start": "2026-10-19", "duration": 7, "focus": "Phase 2 Start", "deliverables": "Healthy Food listing, Food detail, Food cart", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-15", "week": 19, "start": "2026-10-26", "duration": 7, "focus": "Food Complete", "deliverables": "Food order, Order tracking", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-16", "week": 20, "start": "2026-11-02", "duration": 7, "focus": "Shop", "deliverables": "Shop listing, Product detail, Shop cart", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-17", "week": 21, "start": "2026-11-09", "duration": 7, "focus": "Shop Complete", "deliverables": "Shop order, Order tracking", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-18", "week": 22, "start": "2026-11-16", "duration": 7, "focus": "Treatment & Journey", "deliverables": "Treatment transaction, Skin journey", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "BUF-3", "week": 23, "start": "2026-11-23", "duration": 7, "focus": "Buffer", "deliverables": "Phase 2 integration testing, Cross-feature validation, Bug fixes", "team": "Mobile", "color": "#70AD47"},
    {"sprint": "M-19", "week": 24, "start": "2026-11-30", "duration": 7, "focus": "Polish (Part 1)", "deliverables": "UI polish, Animation, Performance", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-20", "week": 25, "start": "2026-12-07", "duration": 7, "focus": "Polish (Part 2)", "deliverables": "Bug fixes, Edge cases, Accessibility", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "M-21", "week": 26, "start": "2026-12-14", "duration": 7, "focus": "QA & Testing", "deliverables": "Integration testing, E2E testing", "team": "Mobile", "color": "#ED7D31"},
    {"sprint": "BUF-4", "week": 27, "start": "2026-12-21", "duration": 7, "focus": "Buffer", "deliverables": "Final QA, Regression testing, App store prep validation", "team": "Mobile", "color": "#70AD47"},
    {"sprint": "M-22", "week": 28, "start": "2026-12-28", "duration": 7, "focus": "Store Prep", "deliverables": "App store assets, Submission, Release", "team": "Mobile", "color": "#ED7D31"},
]

# Milestones
MILESTONES = [
    {"name": "Backend APIs Ready", "date": "2026-07-20", "color": "#4472C4"},
    {"name": "Auth Complete", "date": "2026-07-20", "color": "#ED7D31"},
    {"name": "Homepage Live", "date": "2026-08-10", "color": "#ED7D31"},
    {"name": "Booking + Payment Live", "date": "2026-09-14", "color": "#ED7D31"},
    {"name": "Profile + Promotions Live", "date": "2026-10-19", "color": "#ED7D31"},
    {"name": "Phase 2 (Food, Shop) Live", "date": "2026-11-23", "color": "#ED7D31"},
    {"name": "App Store Submission", "date": "2027-01-04", "color": "#FF0000"},
]


def create_docx():
    """Create DOCX document from roadmap"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Dermaesthetic Super Apps - MVP Development Roadmap', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f'Date: 2026-06-17')
    doc.add_paragraph(f'Status: Draft - Pending Approval')
    doc.add_paragraph(f'Approach: Phased Parallel Streams (Approach C)')
    doc.add_paragraph(f'Deferred: Laundry, Loyalty & Membership')
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    
    data = [
        ('Metric', 'Value'),
        ('Total Features', '14 (11 Phase 1 + 3 Phase 2)'),
        ('Total Sprints', '28 weeks (1 sprint = 1 week)'),
        ('Buffer Weeks', '4 weeks mobile + 2 weeks backend'),
        ('Team', '2 Flutter devs, 2 Fullstack (CMS/BE)'),
        ('Timeline', 'June 2026 → January 2027'),
        ('CTO Target', 'Beat by ~5 months'),
        ('Go-Live', 'January 2027'),
    ]
    
    for i, (key, value) in enumerate(data):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    doc.add_paragraph()
    
    # Buffer Strategy
    doc.add_heading('Buffer Strategy', 2)
    doc.add_paragraph('Buffer weeks are strategically placed after critical features to handle integration issues, bug fixes, and testing.')
    
    buffer_table = doc.add_table(rows=7, cols=3)
    buffer_table.style = 'Light Grid Accent 1'
    
    buffer_data = [
        ('Buffer', 'Week', 'Purpose'),
        ('BUF-1', '5', 'Auth integration testing, Bug fixes after critical auth flow'),
        ('BUF-2', '13', 'Booking integration testing, Payment flow validation'),
        ('BUF-3', '23', 'Phase 2 integration testing, Cross-feature validation'),
        ('BUF-4', '27', 'Final QA, Regression testing, App store prep validation'),
        ('B-BUF-1', '9', 'Backend integration testing, Bug fixes, API documentation'),
        ('B-BUF-2', '10', 'Backend performance optimization, Security audit'),
    ]
    
    for i, (col1, col2, col3) in enumerate(buffer_data):
        buffer_table.rows[i].cells[0].text = col1
        buffer_table.rows[i].cells[1].text = col2
        buffer_table.rows[i].cells[2].text = col3
        if i == 0:
            for cell in buffer_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    doc.add_paragraph()
    
    # Buffer Usage Rules
    doc.add_heading('Buffer Usage Rules', 3)
    rules = [
        'Buffer weeks are for integration testing and bug fixes, not new features',
        'Buffer weeks can be used for documentation and code reviews',
        'Buffer weeks can be used for performance optimization',
        'Buffer weeks cannot be used for scope expansion',
        'If no issues arise, buffer weeks can be used for technical debt reduction',
    ]
    for rule in rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_page_break()
    
    # Team Allocation
    doc.add_heading('Team Allocation', 1)
    
    team_table = doc.add_table(rows=5, cols=3)
    team_table.style = 'Light Grid Accent 1'
    
    team_data = [
        ('Role', 'Count', 'Responsibility'),
        ('Flutter Developer 1', '1', 'Mobile app (Android + iOS)'),
        ('Flutter Developer 2', '1', 'Mobile app (Android + iOS)'),
        ('Fullstack Developer 1', '1', 'Backend APIs + CMS'),
        ('Fullstack Developer 2', '1', 'Backend APIs + CMS'),
    ]
    
    for i, (col1, col2, col3) in enumerate(team_data):
        team_table.rows[i].cells[0].text = col1
        team_table.rows[i].cells[1].text = col2
        team_table.rows[i].cells[2].text = col3
        if i == 0:
            for cell in team_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Critical Path: Mobile development (2 Flutter devs)', style='Intense Quote')
    doc.add_paragraph('Parallel Work: Backend APIs can be built 1-2 weeks ahead of mobile consumption', style='Intense Quote')
    
    doc.add_page_break()
    
    # Weekly Sprint Plan - Backend
    doc.add_heading('Weekly Sprint Plan', 1)
    
    doc.add_heading('Backend Sprints (Fullstack Devs)', 2)
    
    backend_sprints = [s for s in SPRINTS if s['team'] == 'Backend']
    
    backend_table = doc.add_table(rows=len(backend_sprints) + 1, cols=4)
    backend_table.style = 'Light Grid Accent 1'
    
    backend_headers = ['Sprint', 'Week', 'Focus', 'Deliverables']
    for j, header in enumerate(backend_headers):
        backend_table.rows[0].cells[j].text = header
        for paragraph in backend_table.rows[0].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    for i, sprint in enumerate(backend_sprints, 1):
        backend_table.rows[i].cells[0].text = sprint['sprint']
        backend_table.rows[i].cells[1].text = str(sprint['week'])
        backend_table.rows[i].cells[2].text = sprint['focus']
        backend_table.rows[i].cells[3].text = sprint['deliverables']
    
    doc.add_paragraph()
    
    # Weekly Sprint Plan - Mobile
    doc.add_heading('Mobile Sprints (Flutter Devs)', 2)
    
    mobile_sprints = [s for s in SPRINTS if s['team'] == 'Mobile']
    
    mobile_table = doc.add_table(rows=len(mobile_sprints) + 1, cols=4)
    mobile_table.style = 'Light Grid Accent 1'
    
    mobile_headers = ['Sprint', 'Week', 'Focus', 'Deliverables']
    for j, header in enumerate(mobile_headers):
        mobile_table.rows[0].cells[j].text = header
        for paragraph in mobile_table.rows[0].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    for i, sprint in enumerate(mobile_sprints, 1):
        mobile_table.rows[i].cells[0].text = sprint['sprint']
        mobile_table.rows[i].cells[1].text = str(sprint['week'])
        mobile_table.rows[i].cells[2].text = sprint['focus']
        mobile_table.rows[i].cells[3].text = sprint['deliverables']
    
    doc.add_page_break()
    
    # Milestones
    doc.add_heading('Milestones', 1)
    
    milestone_table = doc.add_table(rows=len(MILESTONES) + 1, cols=2)
    milestone_table.style = 'Light Grid Accent 1'
    
    milestone_headers = ['Milestone', 'Date']
    for j, header in enumerate(milestone_headers):
        milestone_table.rows[0].cells[j].text = header
        for paragraph in milestone_table.rows[0].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    for i, milestone in enumerate(MILESTONES, 1):
        milestone_table.rows[i].cells[0].text = milestone['name']
        milestone_table.rows[i].cells[1].text = milestone['date']
    
    doc.add_paragraph()
    
    # Risk Register
    doc.add_heading('Risk Register', 1)
    
    risks = [
        ('Flutter dev bottleneck', 'High', 'Backend APIs ready 1-2 weeks ahead, parallel development'),
        ('Doku payment integration issues', 'Medium', 'Early integration testing, Fallback payment method'),
        ('JNE delivery integration issues', 'Medium', 'Early API testing, Mock for development'),
        ('App store rejection', 'Medium', 'Follow guidelines, Pre-submission review'),
        ('Scope creep', 'High', 'Strict MVP scope, Deferred features documented'),
        ('Team burnout', 'Medium', 'Realistic sprint loads, Buffer weeks included'),
        ('Buffer week misuse', 'Medium', 'Buffer weeks are for integration testing and bug fixes, not new features'),
        ('Integration issues', 'High', 'Buffer weeks placed after critical features (Auth, Booking, Phase 2)'),
        ('Timeline slippage', 'Medium', '4 buffer weeks provide 17% schedule flexibility'),
        ('Backend delays', 'Medium', 'Backend has 2 buffer weeks for integration and security testing'),
    ]
    
    risk_table = doc.add_table(rows=len(risks) + 1, cols=3)
    risk_table.style = 'Light Grid Accent 1'
    
    risk_headers = ['Risk', 'Impact', 'Mitigation']
    for j, header in enumerate(risk_headers):
        risk_table.rows[0].cells[j].text = header
        for paragraph in risk_table.rows[0].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    for i, (risk, impact, mitigation) in enumerate(risks, 1):
        risk_table.rows[i].cells[0].text = risk
        risk_table.rows[i].cells[1].text = impact
        risk_table.rows[i].cells[2].text = mitigation
    
    # Save DOCX
    docx_path = os.path.join(OUTPUT_DIR, 'dermaesthetic-mvp-roadmap.docx')
    doc.save(docx_path)
    print(f'DOCX saved: {docx_path}')
    return docx_path


def create_pdf():
    """Create PDF from DOCX using weasyprint"""
    from weasyprint import HTML
    
    # Read the markdown file
    md_path = os.path.join(OUTPUT_DIR, '2026-06-17-dermaesthetic-mvp-roadmap.md')
    with open(md_path, 'r') as f:
        md_content = f.read()
    
    # Convert markdown to HTML (simple conversion)
    import markdown
    html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    
    # Add CSS styling
    css = """
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; font-size: 12px; }
        h1 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }
        h2 { color: #34495e; border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; }
        h3 { color: #7f8c8d; }
        table { border-collapse: collapse; width: 100%; margin: 20px 0; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        th { background-color: #3498db; color: white; }
        tr:nth-child(even) { background-color: #f2f2f2; }
        code { background-color: #f4f4f4; padding: 2px 5px; border-radius: 3px; }
        pre { background-color: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }
        blockquote { border-left: 4px solid #3498db; margin: 0; padding: 10px 20px; background-color: #f9f9f9; }
        ul, ol { margin: 10px 0; padding-left: 30px; }
        li { margin: 5px 0; }
        .mermaid { text-align: center; margin: 20px 0; }
    </style>
    """
    
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        {css}
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    
    # Save HTML for debugging
    html_path = os.path.join(OUTPUT_DIR, 'dermaesthetic-mvp-roadmap.html')
    with open(html_path, 'w') as f:
        f.write(full_html)
    
    # Generate PDF
    pdf_path = os.path.join(OUTPUT_DIR, 'dermaesthetic-mvp-roadmap.pdf')
    HTML(string=full_html).write_pdf(pdf_path)
    print(f'PDF saved: {pdf_path}')
    return pdf_path


def create_excel_gantt():
    """Create color-coded Gantt chart in Excel"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "MVP Gantt Chart"
    
    # Define colors (aRGB format with alpha channel)
    COLORS = {
        'Backend': 'FF4472C4',      # Blue
        'Mobile': 'FFED7D31',       # Orange
        'Buffer': 'FF70AD47',       # Green
        'Milestone': 'FFFF0000',    # Red
        'Header': 'FF2F5496',       # Dark Blue
        'Weekend': 'FFD9E2F3',      # Light Blue
    }
    
    # Title
    ws.merge_cells('A1:Z1')
    ws['A1'] = 'Dermaesthetic Super Apps - MVP Gantt Chart'
    ws['A1'].font = Font(name='Calibri', size=16, bold=True, color='FFFFFF')
    ws['A1'].fill = PatternFill(start_color=COLORS['Header'], end_color=COLORS['Header'], fill_type='solid')
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 30
    
    # Subtitle
    ws.merge_cells('A2:Z2')
    ws['A2'] = 'Timeline: June 2026 - January 2027 | Team: 2 Flutter Devs + 2 Fullstack Devs'
    ws['A2'].font = Font(name='Calibri', size=11, italic=True, color='FFFFFF')
    ws['A2'].fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    ws['A2'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[2].height = 25
    
    # Headers
    headers = ['Sprint', 'Team', 'Week', 'Focus', 'Deliverables', 'Start Date', 'End Date']
    
    # Calculate number of weeks
    start_date = datetime(2026, 6, 22)
    end_date = datetime(2027, 1, 4)
    num_weeks = ((end_date - start_date).days // 7) + 1
    
    # Add week columns
    for week in range(num_weeks):
        week_start = start_date + timedelta(weeks=week)
        headers.append(f'W{week + 1}\n{week_start.strftime("%b %d")}')
    
    # Write headers
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=header)
        cell.font = Font(name='Calibri', size=10, bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color=COLORS['Header'], end_color=COLORS['Header'], fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = Border(
            left=Side(style='thin', color='FFFFFF'),
            right=Side(style='thin', color='FFFFFF'),
            top=Side(style='thin', color='FFFFFF'),
            bottom=Side(style='thin', color='FFFFFF')
        )
    
    ws.row_dimensions[4].height = 40
    
    # Set column widths
    ws.column_dimensions['A'].width = 12  # Sprint
    ws.column_dimensions['B'].width = 10  # Team
    ws.column_dimensions['C'].width = 8   # Week
    ws.column_dimensions['D'].width = 20  # Focus
    ws.column_dimensions['E'].width = 40  # Deliverables
    ws.column_dimensions['F'].width = 12  # Start Date
    ws.column_dimensions['G'].width = 12  # End Date
    
    for week in range(num_weeks):
        col = get_column_letter(8 + week)
        ws.column_dimensions[col].width = 8
    
    # Write sprint data
    row = 5
    for sprint in SPRINTS:
        sprint_start = datetime.strptime(sprint['start'], '%Y-%m-%d')
        sprint_end = sprint_start + timedelta(days=sprint['duration'] - 1)
        
        # Calculate week number for Gantt bar
        week_num = ((sprint_start - start_date).days // 7) + 1
        duration_weeks = (sprint['duration'] + 6) // 7  # Round up
        
        # Sprint info
        ws.cell(row=row, column=1, value=sprint['sprint']).font = Font(name='Calibri', size=10, bold=True)
        ws.cell(row=row, column=2, value=sprint['team']).font = Font(name='Calibri', size=10)
        ws.cell(row=row, column=3, value=sprint['week']).font = Font(name='Calibri', size=10)
        ws.cell(row=row, column=4, value=sprint['focus']).font = Font(name='Calibri', size=10)
        ws.cell(row=row, column=5, value=sprint['deliverables']).font = Font(name='Calibri', size=9)
        ws.cell(row=row, column=6, value=sprint_start.strftime('%Y-%m-%d')).font = Font(name='Calibri', size=9)
        ws.cell(row=row, column=7, value=sprint_end.strftime('%Y-%m-%d')).font = Font(name='Calibri', size=9)
        
        # Apply row styling
        for col in range(1, 8):
            cell = ws.cell(row=row, column=col)
            cell.alignment = Alignment(vertical='center', wrap_text=True)
            cell.border = Border(
                left=Side(style='thin', color='D9D9D9'),
                right=Side(style='thin', color='D9D9D9'),
                top=Side(style='thin', color='D9D9D9'),
                bottom=Side(style='thin', color='D9D9D9')
            )
        
        # Draw Gantt bar
        bar_color = 'FF' + sprint['color'].lstrip('#')  # Add alpha prefix and remove # for aRGB format
        for week_offset in range(duration_weeks):
            col = 7 + week_num + week_offset
            if col <= len(headers):
                cell = ws.cell(row=row, column=col)
                cell.fill = PatternFill(start_color=bar_color, end_color=bar_color, fill_type='solid')
                cell.border = Border(
                    left=Side(style='thin', color='FFFFFF'),
                    right=Side(style='thin', color='FFFFFF'),
                    top=Side(style='thin', color='FFFFFF'),
                    bottom=Side(style='thin', color='FFFFFF')
                )
        
        row += 1
    
    # Add milestones section
    row += 2
    ws.merge_cells(f'A{row}:G{row}')
    ws.cell(row=row, column=1, value='MILESTONES').font = Font(name='Calibri', size=12, bold=True, color='FFFFFF')
    ws.cell(row=row, column=1).fill = PatternFill(start_color=COLORS['Milestone'], end_color=COLORS['Milestone'], fill_type='solid')
    ws.cell(row=row, column=1).alignment = Alignment(horizontal='center')
    
    row += 1
    for milestone in MILESTONES:
        milestone_date = datetime.strptime(milestone['date'], '%Y-%m-%d')
        week_num = ((milestone_date - start_date).days // 7) + 1
        
        ws.cell(row=row, column=1, value='◆').font = Font(name='Calibri', size=12, color=COLORS['Milestone'])
        ws.cell(row=row, column=2, value=milestone['name']).font = Font(name='Calibri', size=10, bold=True)
        ws.cell(row=row, column=3, value=milestone['date']).font = Font(name='Calibri', size=10)
        
        # Mark milestone on timeline
        col = 7 + week_num
        if col <= len(headers):
            cell = ws.cell(row=row, column=col)
            cell.value = '◆'
            cell.font = Font(name='Calibri', size=12, color=COLORS['Milestone'])
            cell.alignment = Alignment(horizontal='center')
        
        row += 1
    
    # Add legend
    row += 2
    ws.merge_cells(f'A{row}:G{row}')
    ws.cell(row=row, column=1, value='LEGEND').font = Font(name='Calibri', size=12, bold=True)
    
    row += 1
    legend_items = [
        ('Backend Development', COLORS['Backend']),
        ('Mobile Development', COLORS['Mobile']),
        ('Buffer Week', COLORS['Buffer']),
        ('Milestone', COLORS['Milestone']),
    ]
    
    for label, color in legend_items:
        ws.cell(row=row, column=1, value='█').font = Font(name='Calibri', size=12, color=color[2:])  # Remove alpha for font color
        ws.cell(row=row, column=2, value=label).font = Font(name='Calibri', size=10)
        row += 1
    
    # Freeze panes
    ws.freeze_panes = 'H5'
    
    # Save Excel
    excel_path = os.path.join(OUTPUT_DIR, 'dermaesthetic-mvp-gantt.xlsx')
    wb.save(excel_path)
    print(f'Excel Gantt saved: {excel_path}')
    return excel_path


if __name__ == '__main__':
    print('Generating MVP roadmap documents...')
    print()
    
    # Create DOCX
    print('1. Creating DOCX...')
    docx_path = create_docx()
    
    # Create PDF
    print('2. Creating PDF...')
    pdf_path = create_pdf()
    
    # Create Excel Gantt
    print('3. Creating Excel Gantt chart...')
    excel_path = create_excel_gantt()
    
    print()
    print('All documents generated successfully!')
    print(f'  - DOCX: {docx_path}')
    print(f'  - PDF: {pdf_path}')
    print(f'  - Excel: {excel_path}')
